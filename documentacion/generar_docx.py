"""
Generador de muestra-visual.docx — Refugio del Sol.

Reproduce, dentro de las posibilidades de Word, la maqueta visual del archivo
muestra-visual.html: portada, índice y sección "Finalidad del proyecto" con
encabezado, pie de página, paleta cromática y tipografía Poppins.

Uso:
    python generar_docx.py

Requiere:
    pip install python-docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# =============================================================================
# Paleta y constantes
# =============================================================================

C_MOSTAZA       = RGBColor(0xE9, 0xC4, 0x6A)
C_MOSTAZA_OSC   = RGBColor(0xD4, 0xA9, 0x43)
C_SAGE          = RGBColor(0x80, 0xAB, 0x9F)
C_SAGE_OSC      = RGBColor(0x6C, 0x92, 0x86)
C_BERENJENA     = RGBColor(0x4F, 0x3D, 0x63)
C_BERENJENA_OSC = RGBColor(0x2D, 0x1F, 0x3D)
C_CREMA         = "F2E9E4"   # como hex string para shading
C_TEXTO         = RGBColor(0x2D, 0x1F, 0x3D)
C_TEXTO_SUAVE   = RGBColor(0x6B, 0x5B, 0x73)

FUENTE = "Poppins"  # con fallback automático a Calibri si no está instalada

BASE = Path(__file__).resolve().parent
LOGO = BASE / "assets" / "logo.png"
SALIDA = BASE / "muestra-visual.docx"


# =============================================================================
# Helpers de bajo nivel (XML para detalles que python-docx no expone)
# =============================================================================

def set_cell_background(cell, hex_color):
    """Pinta el fondo de una celda con un color hex sin '#'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_left_border(cell, hex_color, size_pt=18):
    """Añade un borde izquierdo grueso a una celda (para callouts)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt))   # eighths of a point
    left.set(qn("w:color"), hex_color)
    borders.append(left)
    tc_pr.append(borders)


def remove_cell_borders(cell):
    """Quita todos los bordes visibles de una celda salvo los que se añadan."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)


def add_horizontal_line(paragraph, hex_color="80AB9F", size_pt=4):
    """Inserta una línea horizontal usando borde inferior del párrafo."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_top_line(paragraph, hex_color="E9C46A", size_pt=4):
    """Línea superior en un párrafo (para header/footer)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_pt))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), hex_color)
    p_bdr.append(top)
    p_pr.append(p_bdr)


def add_field(paragraph, instr_text, font_name=FUENTE, font_size=9,
              color=None, italic=False):
    """Inserta un campo de Word (p.ej. PAGE) en un párrafo."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def styled_run(paragraph, text, *, size=11, bold=False, italic=False,
               color=C_TEXTO, font=FUENTE, letter_spacing=None,
               uppercase=False):
    """Añade un run con estilo tipográfico fino."""
    run = paragraph.add_run(text.upper() if uppercase else text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # asegura la fuente para Asian fallback (Word la respeta mejor)
    rPr = run._r.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    rFonts = existing if existing is not None else OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    if rFonts.getparent() is None:
        rPr.append(rFonts)
    if letter_spacing is not None:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(letter_spacing * 20)))  # in twentieths of a point
        rPr.append(spacing)
    return run


def set_section_margins(section, top=22, bottom=22, left=22, right=22):
    """Márgenes en milímetros."""
    section.top_margin = Mm(top)
    section.bottom_margin = Mm(bottom)
    section.left_margin = Mm(left)
    section.right_margin = Mm(right)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)


def set_section_a4(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)


# =============================================================================
# Construcción del documento
# =============================================================================

def build():
    doc = Document()

    # ---------- Estilos base globales ----------
    style_normal = doc.styles["Normal"]
    style_normal.font.name = FUENTE
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = C_TEXTO
    rPr = style_normal.element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    rFonts = existing if existing is not None else OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FUENTE)
    rFonts.set(qn("w:hAnsi"), FUENTE)
    if rFonts.getparent() is None:
        rPr.append(rFonts)

    # =========================================================================
    # SECCIÓN 1 — PORTADA (sin header / footer)
    # =========================================================================
    section1 = doc.sections[0]
    set_section_a4(section1)
    set_section_margins(section1, top=0, bottom=0, left=0, right=0)
    section1.different_first_page_header_footer = False

    # Banda superior mostaza (tabla 1×1 con shading)
    band = doc.add_table(rows=1, cols=1)
    band.autofit = False
    band.columns[0].width = Mm(210)
    band_cell = band.rows[0].cells[0]
    band_cell.width = Mm(210)
    set_cell_background(band_cell, "E9C46A")
    remove_cell_borders(band_cell)
    band_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    band_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    # altura mínima
    tr = band.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "454")  # ~8mm
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    # Espacio superior antes del eyebrow
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(0)
    p_spacer.paragraph_format.left_indent = Mm(28)
    p_spacer.paragraph_format.right_indent = Mm(28)
    styled_run(p_spacer, "", size=1)
    p_spacer.paragraph_format.space_after = Mm(20)

    # Eyebrow
    p_eyebrow = doc.add_paragraph()
    p_eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eyebrow.paragraph_format.space_after = Mm(14)
    styled_run(p_eyebrow, "PROYECTO FINAL · 2º DAW",
               size=9, color=C_SAGE_OSC, letter_spacing=4, bold=True)

    # Logo centrado
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Mm(12)
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO), width=Mm(60))
    else:
        print(f"[aviso] Logo no encontrado en {LOGO}")

    # Título técnico
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Mm(8)
    p_title.paragraph_format.left_indent = Mm(28)
    p_title.paragraph_format.right_indent = Mm(28)
    styled_run(p_title, "Plataforma Web para la Gestión Integral",
               size=26, bold=True, color=C_BERENJENA)
    p_title.add_run().add_break()
    styled_run(p_title, "de Asociaciones Felinas",
               size=26, bold=True, color=C_BERENJENA)

    # Línea sage corta
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_after = Mm(8)
    styled_run(p_rule, " " * 20, size=1)
    add_horizontal_line(p_rule, hex_color="80AB9F", size_pt=8)

    # Marca
    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_brand.paragraph_format.space_after = Mm(3)
    styled_run(p_brand, "Refugio del Sol", size=22, color=C_MOSTAZA_OSC)

    # Tagline
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_after = Mm(20)
    styled_run(p_tag, "“Que el sol de su vida sea un hogar”",
               size=14, italic=True, color=C_BERENJENA, font="Cambria")

    # Bloque meta (autora, curso, centro) — tabla 2×4 con borde izquierdo
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.autofit = False
    meta_table.columns[0].width = Mm(45)
    meta_table.columns[1].width = Mm(75)
    meta_data = [
        ("AUTORA", "Viviana Salazar"),
        ("CURSO", "2º Desarrollo de Aplicaciones Web"),
        ("CENTRO", "IES Belén"),
        ("CURSO ACADÉMICO", "2025 / 2026"),
    ]
    for i, (k, v) in enumerate(meta_data):
        c0 = meta_table.rows[i].cells[0]
        c0.width = Mm(45)
        c1 = meta_table.rows[i].cells[1]
        c1.width = Mm(75)
        remove_cell_borders(c0)
        remove_cell_borders(c1)
        if i == 0:
            set_cell_left_border(c0, "E9C46A", size_pt=20)
        else:
            set_cell_left_border(c0, "E9C46A", size_pt=20)
        # eyebrow
        p_k = c0.paragraphs[0]
        p_k.paragraph_format.left_indent = Mm(4)
        styled_run(p_k, k, size=8, bold=True, color=C_SAGE_OSC, letter_spacing=2)
        # valor
        p_v = c1.paragraphs[0]
        styled_run(p_v, v, size=11, bold=True, color=C_BERENJENA)

    # alineamos la tabla a la izquierda con sangría
    tbl_pr = meta_table._tbl.tblPr
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(int(28 * 56.7)))  # 28mm
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    # Espacio antes de la banda inferior
    p_filler = doc.add_paragraph()
    p_filler.paragraph_format.space_after = Mm(40)
    styled_run(p_filler, "", size=1)

    # Banda inferior sage
    band2 = doc.add_table(rows=1, cols=1)
    band2.autofit = False
    band2.columns[0].width = Mm(210)
    band2_cell = band2.rows[0].cells[0]
    set_cell_background(band2_cell, "80AB9F")
    remove_cell_borders(band2_cell)
    tr2 = band2.rows[0]._tr
    trPr2 = tr2.get_or_add_trPr()
    trHeight2 = OxmlElement("w:trHeight")
    trHeight2.set(qn("w:val"), "284")  # ~5mm
    trHeight2.set(qn("w:hRule"), "exact")
    trPr2.append(trHeight2)

    # =========================================================================
    # SECCIÓN 2 — RESTO DEL DOCUMENTO (con header / footer)
    # =========================================================================
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_a4(section2)
    set_section_margins(section2, top=24, bottom=22, left=22, right=22)
    section2.different_first_page_header_footer = False

    # romper la herencia del header/footer respecto a la sección 1
    section2.header.is_linked_to_previous = False
    section2.footer.is_linked_to_previous = False

    # ---------- Header ----------
    header_table = section2.header.add_table(rows=1, cols=2, width=Mm(166))
    header_table.autofit = False
    header_table.columns[0].width = Mm(83)
    header_table.columns[1].width = Mm(83)
    hl, hr = header_table.rows[0].cells
    remove_cell_borders(hl)
    remove_cell_borders(hr)
    p_hl = hl.paragraphs[0]
    styled_run(p_hl, "REFUGIO DEL SOL", size=9, color=C_BERENJENA, letter_spacing=1.2, bold=False)
    add_horizontal_line(p_hl, hex_color="E9C46A", size_pt=4)
    p_hr = hr.paragraphs[0]
    p_hr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    styled_run(p_hr, "Documentación del Proyecto", size=9, color=C_SAGE)
    add_horizontal_line(p_hr, hex_color="E9C46A", size_pt=4)

    # ---------- Footer ----------
    footer_table = section2.footer.add_table(rows=1, cols=2, width=Mm(166))
    footer_table.autofit = False
    footer_table.columns[0].width = Mm(83)
    footer_table.columns[1].width = Mm(83)
    fl, fr = footer_table.rows[0].cells
    remove_cell_borders(fl)
    remove_cell_borders(fr)
    p_fl = fl.paragraphs[0]
    add_top_line(p_fl, hex_color="80AB9F", size_pt=4)
    styled_run(p_fl, "Viviana Salazar  ·  2º DAW", size=9, color=C_BERENJENA)
    p_fr = fr.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_top_line(p_fr, hex_color="80AB9F", size_pt=4)
    styled_run(p_fr, "— ", size=9, color=C_BERENJENA, italic=True)
    add_field(p_fr, "PAGE", color=C_BERENJENA, italic=True)
    styled_run(p_fr, " —", size=9, color=C_BERENJENA, italic=True)

    # =========================================================================
    # ÍNDICE
    # =========================================================================
    p_eyeb = doc.add_paragraph()
    p_eyeb.paragraph_format.space_before = Mm(4)
    p_eyeb.paragraph_format.space_after = Mm(2)
    styled_run(p_eyeb, "TABLA DE CONTENIDOS", size=9, color=C_SAGE_OSC,
               bold=True, letter_spacing=4)

    p_titulo_idx = doc.add_paragraph()
    p_titulo_idx.paragraph_format.space_after = Mm(2)
    styled_run(p_titulo_idx, "Índice", size=34, bold=True, color=C_BERENJENA)

    p_idx_rule = doc.add_paragraph()
    p_idx_rule.paragraph_format.space_after = Mm(10)
    styled_run(p_idx_rule, " " * 12, size=1)
    add_horizontal_line(p_idx_rule, hex_color="E9C46A", size_pt=12)

    INDICE = [
        ("01", "Título del proyecto", "01", []),
        ("02", "Índice", "02", []),
        ("03", "Estudio del sector productivo", "04", []),
        ("04", "Finalidad del proyecto", "07", []),
        ("05", "Descripción básica del proyecto", "11", []),
        ("06", "Descripción detallada del proyecto", "15", [
            ("6.1",  "Herramientas de desarrollo software utilizadas", "16"),
            ("6.2",  "Diseño del proyecto", "19"),
            ("6.3",  "Esquema de funcionamiento de la aplicación", "23"),
            ("6.4",  "Descomposición modular e interrelación entre módulos", "27"),
            ("6.5",  "Descripción de los módulos del proyecto", "31"),
            ("6.6",  "Descripción de la interfaz de usuario", "38"),
            ("6.7",  "Descripción de listados e informes", "44"),
            ("6.8",  "Manual de usuario", "47"),
            ("6.9",  "Manual de instalación o despliegue", "53"),
            ("6.10", "Presentación de la aplicación", "58"),
        ]),
        ("07", "Planificación del proyecto", "61", []),
        ("08", "Control de la ejecución", "65", []),
        ("09", "Dificultades encontradas", "68", []),
        ("10", "Propuestas de mejora", "71", []),
        ("11", "Conclusión final", "74", []),
    ]

    for num, title, page, subs in INDICE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # tab stop derecho con leader de puntos a 166mm
        p.paragraph_format.tab_stops.add_tab_stop(
            Mm(166), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        styled_run(p, f"{num}    ", size=11, bold=True, color=C_MOSTAZA_OSC)
        styled_run(p, title, size=11, color=C_BERENJENA, bold=True)
        styled_run(p, "\t", size=11)
        styled_run(p, page, size=10, color=C_TEXTO_SUAVE)

        for sn, st, sp in subs:
            ps = doc.add_paragraph()
            ps.paragraph_format.left_indent = Mm(14)
            ps.paragraph_format.space_before = Pt(0)
            ps.paragraph_format.space_after = Pt(0)
            ps.paragraph_format.tab_stops.add_tab_stop(
                Mm(152), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
            )
            styled_run(ps, f"{sn}    ", size=10, color=C_SAGE)
            styled_run(ps, st, size=10, color=C_SAGE_OSC)
            styled_run(ps, "\t", size=10)
            styled_run(ps, sp, size=10, color=C_TEXTO_SUAVE)

    # Salto de página antes de la sección de ejemplo
    doc.add_page_break()

    # =========================================================================
    # SECCIÓN: FINALIDAD DEL PROYECTO
    # =========================================================================
    p_eyeb2 = doc.add_paragraph()
    p_eyeb2.paragraph_format.space_before = Mm(4)
    p_eyeb2.paragraph_format.space_after = Mm(2)
    styled_run(p_eyeb2, "SECCIÓN 04", size=9, color=C_SAGE_OSC, bold=True,
               letter_spacing=4)

    p_sec_title = doc.add_paragraph()
    p_sec_title.paragraph_format.space_after = Mm(2)
    styled_run(p_sec_title, "Finalidad del proyecto",
               size=28, bold=True, color=C_BERENJENA)

    p_sec_rule = doc.add_paragraph()
    p_sec_rule.paragraph_format.space_after = Mm(8)
    styled_run(p_sec_rule, " " * 10, size=1)
    add_horizontal_line(p_sec_rule, hex_color="80AB9F", size_pt=12)

    # Lead
    p_lead = doc.add_paragraph()
    p_lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_lead.paragraph_format.space_after = Mm(6)
    styled_run(p_lead,
               "El objetivo primordial de ", size=12.5, color=C_BERENJENA, bold=False)
    styled_run(p_lead, "Refugio del Sol", size=12.5, bold=True, color=C_BERENJENA)
    styled_run(p_lead,
               " es dotar a una asociación de rescate y acogida felina de una "
               "herramienta digital centralizada y eficiente que facilite su "
               "administración diaria, amplíe su alcance hacia adoptantes, "
               "padrinos y voluntarios, y fortalezca el vínculo emocional y "
               "financiero con la comunidad existente.",
               size=12.5, color=C_BERENJENA)

    # H2
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Mm(6)
    p_h2.paragraph_format.space_after = Mm(3)
    styled_run(p_h2, "▍  ", size=14, color=C_MOSTAZA, bold=True)
    styled_run(p_h2, "Necesidad detectada en el sector",
               size=15, color=C_MOSTAZA_OSC, bold=True)

    parrafos_necesidad = [
        "En la mayoría de webs de asociaciones animales actuales la información "
        "se presenta de forma incompleta o difícil de localizar. Las solicitudes "
        "de adopción se reducen a formularios sin respuesta inmediata o a textos "
        "demasiado extensos que terminan por desalentar al usuario interesado.",

        "No existe, en la práctica, un espacio digital en el que adoptantes, "
        "voluntarios y padrinos puedan acceder a información personalizada o "
        "realizar un seguimiento real de sus actividades. Los padrinos, en "
        "concreto, no reciben actualizaciones —fotografías, vídeos o pequeños "
        "informes— de los animales apadrinados, lo que rompe el vínculo "
        "emocional y, con frecuencia, conduce al abandono del apoyo económico.",

        "Las tiendas solidarias y los canales de donación suelen estar alojados "
        "en plataformas externas como Instagram, Vinted o Wallapop. Esto obliga "
        "al usuario a abandonar el sitio principal y a registrarse en otros "
        "portales, reduciendo la probabilidad de compra o donación. A esto se "
        "suma que los gatos en adopción o acogida no siempre están bien "
        "representados: a menudo solo se muestran los casos urgentes, con "
        "información en texto plano y sin estructura visual.",
    ]
    for txt in parrafos_necesidad:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Mm(3)
        p.paragraph_format.line_spacing = 1.55
        styled_run(p, txt, size=11, color=C_TEXTO)

    # Callout — tabla 1×1 con shading crema y borde izquierdo sage
    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Mm(166)
    cell_callout = callout.rows[0].cells[0]
    set_cell_background(cell_callout, C_CREMA)
    remove_cell_borders(cell_callout)
    set_cell_left_border(cell_callout, "80AB9F", size_pt=24)

    p_label = cell_callout.paragraphs[0]
    p_label.paragraph_format.space_after = Mm(2)
    styled_run(p_label, "IDEA FUERZA", size=8.5, bold=True, color=C_SAGE_OSC,
               letter_spacing=2)

    p_callout = cell_callout.add_paragraph()
    p_callout.paragraph_format.space_after = Pt(0)
    styled_run(p_callout,
               "La nueva plataforma se plantea como una ", size=11,
               color=C_BERENJENA)
    styled_run(p_callout, "solución integral", size=11, bold=True,
               color=C_BERENJENA)
    styled_run(p_callout,
               " que centraliza la información, facilita la interacción "
               "personalizada con la comunidad e integra tienda solidaria y "
               "pasarela de pago en un único entorno; abriendo además la puerta "
               "a colaboraciones y afiliaciones con marcas de productos para "
               "mascotas como nueva fuente de ingresos sostenibles.",
               size=11, color=C_BERENJENA)

    # padding del callout
    for p in cell_callout.paragraphs:
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.right_indent = Mm(4)

    # H2 solución
    p_h2b = doc.add_paragraph()
    p_h2b.paragraph_format.space_before = Mm(8)
    p_h2b.paragraph_format.space_after = Mm(3)
    styled_run(p_h2b, "▍  ", size=14, color=C_MOSTAZA, bold=True)
    styled_run(p_h2b, "Solución propuesta",
               size=15, color=C_MOSTAZA_OSC, bold=True)

    p_intro_sol = doc.add_paragraph()
    p_intro_sol.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro_sol.paragraph_format.space_after = Mm(3)
    p_intro_sol.paragraph_format.line_spacing = 1.55
    styled_run(p_intro_sol,
               "La plataforma articula la respuesta a las necesidades "
               "anteriores en torno a cuatro ejes complementarios, pensados "
               "para reforzarse entre sí y simplificar el día a día tanto del "
               "equipo gestor como de los usuarios externos:",
               size=11, color=C_TEXTO)

    bullets = [
        ("Centralización", " de la información y de la gestión interna en un "
         "único panel coherente, accesible desde cualquier dispositivo."),
        ("Interacción personalizada", " con adoptantes, padrinos y "
         "voluntarios, con seguimiento del estado de cada solicitud y un feed "
         "privado para los animales apadrinados."),
        ("Tienda solidaria y pasarela de pago integradas", ", eliminando la "
         "fricción de plataformas externas y aumentando la conversión de la "
         "intención de ayuda en aportación real."),
        ("Logística interna ordenada", " mediante una lista de tareas "
         "compartida y una pizarra comunitaria, que mejora la organización de "
         "traslados, repartos y mantenimiento de los espacios físicos."),
    ]
    for bold_part, rest in bullets:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(8)
        p.paragraph_format.space_after = Mm(2)
        p.paragraph_format.line_spacing = 1.5
        styled_run(p, "■  ", size=10, color=C_MOSTAZA, bold=True)
        styled_run(p, bold_part, size=11, bold=True, color=C_BERENJENA)
        styled_run(p, rest, size=11, color=C_TEXTO)

    p_outro = doc.add_paragraph()
    p_outro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_outro.paragraph_format.space_before = Mm(3)
    p_outro.paragraph_format.line_spacing = 1.55
    styled_run(p_outro,
               "De este modo, la herramienta no se limita a digitalizar un "
               "proceso existente: redefine el modelo de relación entre la "
               "asociación y su comunidad, y prepara a la organización para "
               "escalar su impacto sin multiplicar su carga administrativa.",
               size=11, color=C_TEXTO)

    # Cierre decorativo
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Mm(10)
    styled_run(p_end, " " * 6, size=1)
    add_horizontal_line(p_end, hex_color="80AB9F", size_pt=4)

    # ---------- Guardar ----------
    doc.save(SALIDA)
    print(f"[OK] Documento generado: {SALIDA}")


if __name__ == "__main__":
    build()
