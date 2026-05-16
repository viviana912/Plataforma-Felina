"""
Generador de muestra-visual-2.docx — Refugio del Sol (versión sobria).

Mismo contenido que generar_docx.py con tres cambios:
  1. La marca "Refugio del Sol" y el eslogan quedan claramente centrados
     en el ancho de la página de portada (sin indentaciones residuales).
  2. Se elimina la enumeración de sección que aparecía arriba a la derecha
     y la línea mostaza decorativa del encabezado.
  3. Se elimina la "primera letra grande en mostaza" (lead destacado):
     el primer párrafo se renderiza con el mismo tamaño y color que el
     cuerpo del documento.

Uso:
    python generar_docx_v2.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# Paleta y constantes
# =============================================================================

C_MOSTAZA       = RGBColor(0xE9, 0xC4, 0x6A)
C_MOSTAZA_OSC   = RGBColor(0xD4, 0xA9, 0x43)
C_SAGE          = RGBColor(0x80, 0xAB, 0x9F)
C_SAGE_OSC      = RGBColor(0x6C, 0x92, 0x86)
C_BERENJENA     = RGBColor(0x4F, 0x3D, 0x63)
C_BERENJENA_OSC = RGBColor(0x2D, 0x1F, 0x3D)
C_CREMA         = "F2E9E4"
C_TEXTO         = RGBColor(0x2D, 0x1F, 0x3D)
C_TEXTO_SUAVE   = RGBColor(0x6B, 0x5B, 0x73)

# Colores específicos de la versión sobria
C_LINEA_GRIS    = "d4cad4"   # línea gris suave que sustituye al mostaza

FUENTE = "Poppins"

BASE = Path(__file__).resolve().parent
LOGO = BASE / "assets" / "logo.png"
SALIDA = BASE / "muestra-visual-2.docx"


# =============================================================================
# Helpers
# =============================================================================

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_left_border(cell, hex_color, size_pt=18):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt))
    left.set(qn("w:color"), hex_color)
    borders.append(left)
    tc_pr.append(borders)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tc_pr.append(borders)


def add_horizontal_line(paragraph, hex_color="d4cad4", size_pt=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_top_line(paragraph, hex_color="d4cad4", size_pt=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size_pt))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), hex_color)
    p_bdr.append(top)
    p_pr.append(p_bdr)


def add_field(paragraph, instr_text, font_name=FUENTE, font_size=9,
              color=None, italic=False):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def styled_run(paragraph, text, *, size=11, bold=False, italic=False,
               color=C_TEXTO, font=FUENTE, letter_spacing=None,
               uppercase=False):
    run = paragraph.add_run(text.upper() if uppercase else text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    rFonts = existing if existing is not None else OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    if rFonts.getparent() is None:
        rPr.append(rFonts)
    if letter_spacing is not None:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(letter_spacing * 20)))
        rPr.append(spacing)
    return run


def set_section_margins(section, top=22, bottom=22, left=22, right=22):
    section.top_margin = Mm(top)
    section.bottom_margin = Mm(bottom)
    section.left_margin = Mm(left)
    section.right_margin = Mm(right)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)


def set_section_a4(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)


def reset_horizontal_indent(paragraph):
    """Asegura que un párrafo no tenga indentación lateral residual."""
    paragraph.paragraph_format.left_indent = Mm(0)
    paragraph.paragraph_format.right_indent = Mm(0)
    paragraph.paragraph_format.first_line_indent = Mm(0)


# =============================================================================
# Helpers de alto nivel para construir secciones
# =============================================================================

def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Mm(4)
    p.paragraph_format.space_after = Mm(2)
    styled_run(p, title, size=28, bold=True, color=C_BERENJENA)

    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_after = Mm(8)
    styled_run(p_rule, " " * 10, size=1)
    add_horizontal_line(p_rule, hex_color="80AB9F", size_pt=12)


def add_subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Mm(6)
    p.paragraph_format.space_after = Mm(3)
    styled_run(p, "▍  ", size=14, color=C_MOSTAZA, bold=True)
    styled_run(p, text, size=15, color=C_MOSTAZA_OSC, bold=True)


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Mm(4)
    p.paragraph_format.space_after = Mm(2)
    styled_run(p, text, size=12, color=C_SAGE_OSC, bold=True)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Mm(3)
    p.paragraph_format.line_spacing = 1.55
    styled_run(p, text, size=11, color=C_TEXTO)
    return p


def add_lead(doc, runs):
    """`runs` es una lista de tuplas (texto, bold)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Mm(4)
    p.paragraph_format.line_spacing = 1.7
    for text, bold in runs:
        styled_run(p, text, size=11, color=C_TEXTO, bold=bold)


def add_bullets(doc, items):
    """`items` puede ser str (bullet simple) o tupla (bold_part, rest)."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(8)
        p.paragraph_format.space_after = Mm(2)
        p.paragraph_format.line_spacing = 1.5
        styled_run(p, "■  ", size=10, color=C_MOSTAZA, bold=True)
        if isinstance(item, tuple):
            bold_part, rest = item
            styled_run(p, bold_part, size=11, bold=True, color=C_BERENJENA)
            styled_run(p, rest, size=11, color=C_TEXTO)
        else:
            styled_run(p, item, size=11, color=C_TEXTO)


def add_callout(doc, runs, label="IDEA FUERZA"):
    """Tarjeta con fondo crema y borde lateral sage. `runs` lista de (texto, bold)."""
    callout = doc.add_table(rows=1, cols=1)
    callout.autofit = False
    callout.columns[0].width = Mm(166)
    cell = callout.rows[0].cells[0]
    set_cell_background(cell, C_CREMA)
    remove_cell_borders(cell)
    set_cell_left_border(cell, "80AB9F", size_pt=24)

    if label:
        p_label = cell.paragraphs[0]
        p_label.paragraph_format.space_after = Mm(2)
        styled_run(p_label, label, size=8.5, bold=True, color=C_SAGE_OSC, letter_spacing=2)
        p_text = cell.add_paragraph()
    else:
        p_text = cell.paragraphs[0]
    p_text.paragraph_format.space_after = Pt(0)
    for text, bold in runs:
        styled_run(p_text, text, size=11, bold=bold, color=C_BERENJENA)

    for p in cell.paragraphs:
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.right_indent = Mm(4)


def add_close_rule(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Mm(10)
    styled_run(p, " " * 6, size=1)
    add_horizontal_line(p, hex_color="80AB9F", size_pt=4)


# =============================================================================
# Secciones de contenido
# =============================================================================

def build_seccion_3(doc):
    add_section_title(doc, "03 Estudio del sector productivo")

    add_subhead(doc, "Tipos de empresas del sector")
    add_para(doc,
        "Las empresas y proyectos de necesidades sociales ofrecen productos o servicios "
        "de bienestar, educación y formación. Centrándonos en el bienestar animal, "
        "encontramos las asociaciones felinas de Málaga que dependen al cien por cien "
        "de casas de acogida, como Gatos Sin Hogar.")

    add_subhead(doc, "Necesidades más demandadas del sector")
    add_para(doc,
        "Entre las necesidades más demandadas por las empresas y proyectos del sector "
        "destacan las siguientes:")
    add_bullets(doc, [
        "Falta de fondos.",
        "Ineficiencia logística.",
        "Gestión de voluntarios.",
        "Traslado de suministros.",
        "Falta de información clara y accesible.",
    ])

    add_subhead(doc, "Proyectos para satisfacerlas")
    add_bullets(doc, [
        ("Gamificación y fidelización:",
         " sistema de insignias y reconocimientos para captar al usuario y mantener su compromiso a lo largo del tiempo."),
        ("Gestión interna y logística:",
         " fichas de gatos, canalización de las necesidades de suministros de las casas de acogida a través de la pizarra comunitaria y logística externa con una lista de tareas pendientes para los voluntarios."),
        ("Apadrinamiento interactivo:",
         " zona que conecta al padrino con su gato, mediante actualizaciones periódicas con foto y texto que mantienen la conexión y, con ella, la cuota mensual."),
        ("Tienda solidaria:",
         " productos solidarios e integración con pasarela de pago."),
    ])

    add_subhead(doc, "Oportunidades de negocio previsibles")
    add_bullets(doc, [
        ("Suscripción por capas:",
         " ofrecer la plataforma con diferentes planes de precios según el tamaño de la asociación o la cantidad de animales y usuarios gestionados."),
        ("Servicio de logística centralizada:",
         " transporte y entrega centralizados de suministros, gestionados a través del módulo de acogida e inventario."),
        ("Afiliación con marcas de mascotas:",
         " integrar en la tienda solidaria la venta de piensos, juguetes y accesorios de marcas asociadas."),
        ("Publicidad no intrusiva y patrocinios:",
         " vender espacios publicitarios o patrocinios de marcas relevantes del sector del bienestar animal."),
        ("Modelo Freemium y comisión por éxito:",
         " ofrecer los módulos básicos de gestión logística, fichas felinas y gestión de voluntarios de forma gratuita o a un coste simbólico."),
    ])

    add_subhead(doc, "Informe de prevención de riesgos laborales")

    add_h3(doc, "Riesgos biológicos y sanitarios")
    add_bullets(doc, [
        ("Mordeduras y arañazos:",
         " lesiones físicas por el manejo de gatos estresados, asustados o ferales. Prevención: uso de equipos de protección individual como guantes de manejo grueso y establecimiento de un protocolo de actuación y desinfección inmediato tras una lesión."),
        ("Zoonosis:",
         " contagio de enfermedades transmitidas por animales. Prevención: higiene estricta, protocolos de cuarentena para gatos nuevos o enfermos, uso de EPIs y control veterinario riguroso y periódico."),
        ("Alergias:",
         " reacciones al pelo, saliva o caspa felina, o a productos de limpieza y arena. Prevención: ventilación adecuada de las zonas de gatos y sustitución de productos por opciones hipoalergénicas si fuera necesario."),
    ])

    add_h3(doc, "Riesgos ergonómicos y físicos")
    add_bullets(doc, [
        ("Sobreesfuerzos y posturas forzadas:",
         " lesiones dorsolumbares al levantar sacos de pienso, mover jaulas o agacharse repetidamente. Prevención: formación en técnicas de levantamiento de cargas, uso de equipos de ayuda y rotación de tareas para evitar la repetición constante."),
        ("Caídas al mismo nivel:",
         " tropiezos con suministros, bebederos, comederos o por derrames. Prevención: orden y limpieza exhaustivos, calzado antideslizante adecuado."),
    ])

    add_h3(doc, "Riesgos psicosociales")
    add_bullets(doc, [
        ("Estrés por empatía:",
         " desgaste emocional asociado al trato constante con animales enfermos o maltratados, o derivado de tener que tomar decisiones difíciles. Prevención: apoyo psicológico para los voluntarios con exposición alta."),
    ])

    add_h3(doc, "Riesgos derivados de la logística y los suministros")
    add_bullets(doc, [
        ("Traslado de suministros:",
         " accidentes de tráfico o lesiones al manipular y transportar grandes volúmenes de pienso, arena y jaulas hasta las casas de acogida. Prevención: revisión periódica del estado de los vehículos, asegurar correctamente la carga y establecer límites de peso por persona."),
    ])

    add_h3(doc, "Riesgo de cobertura legal")
    add_bullets(doc, [
        ("Ausencia de seguros o convenios:",
         " falta de cobertura legal a los voluntarios ante accidentes o lesiones ocurridas durante la actividad. Prevención: disponer de un seguro de responsabilidad civil y, fundamentalmente, de un seguro de accidentes para voluntarios que cubra los riesgos inherentes a sus tareas."),
    ])

    add_close_rule(doc)


def build_seccion_4(doc):
    add_section_title(doc, "04 Finalidad del proyecto")

    add_lead(doc, [
        ("El objetivo primordial de Refugio del Sol es dotar a una asociación de "
         "rescate y acogida felina de una herramienta digital centralizada y "
         "eficiente que facilite su administración diaria, amplíe su alcance hacia "
         "adoptantes, padrinos y voluntarios, y fortalezca el vínculo emocional y "
         "financiero con la comunidad existente.", False),
    ])

    add_para(doc,
        "En la mayoría de webs de asociaciones animales actuales la información se "
        "presenta de forma incompleta o difícil de localizar. Las solicitudes de "
        "adopción se reducen a formularios sin respuesta inmediata o a textos "
        "demasiado extensos que terminan por desalentar al usuario interesado.")

    add_para(doc,
        "No existe, en la práctica, un espacio digital en el que adoptantes, "
        "voluntarios y padrinos puedan acceder a información personalizada o realizar "
        "un seguimiento real de sus actividades. Los padrinos, en concreto, no "
        "reciben actualizaciones —fotografías, vídeos o pequeños informes— de los "
        "animales apadrinados, lo que rompe el vínculo emocional y, con frecuencia, "
        "conduce al abandono del apoyo económico.")

    add_para(doc,
        "Las tiendas solidarias y los canales de donación suelen estar alojados en "
        "plataformas externas como Instagram, Vinted o Wallapop. Esto obliga al "
        "usuario a abandonar el sitio principal y a registrarse en otros portales, "
        "reduciendo la probabilidad de compra o donación. A esto se suma que los "
        "gatos en adopción o acogida no siempre están bien representados: a menudo "
        "solo se muestran los casos urgentes, con información en texto plano y sin "
        "estructura visual.")

    add_callout(doc, [
        ("La nueva plataforma se plantea como una ", False),
        ("solución integral", True),
        (" que centraliza la información, facilita la interacción personalizada con "
         "la comunidad e ", False),
        ("integra donaciones y apadrinamientos con pasarela de pago ficticia en un "
         "único entorno", True),
        (".", False),
    ])

    add_para(doc,
        "Asimismo, la funcionalidad de gestión logística y la lista de tareas ayudan "
        "a organizar mejor traslados y repartos, contribuyendo a mantener los "
        "espacios seguros y ordenados, previniendo caídas o sobreesfuerzos derivados "
        "del manejo de suministros.")

    add_close_rule(doc)


def build_seccion_5(doc):
    add_section_title(doc, "05 Descripción básica del proyecto")

    add_lead(doc, [
        ("El proyecto consiste en el desarrollo de una ", False),
        ("plataforma web interactiva", True),
        (" diseñada específicamente para optimizar la gestión integral y aumentar la "
         "sostenibilidad de una asociación de rescate y acogida de gatos. El sistema "
         "transforma digitalmente los procesos internos para dar autonomía y "
         "eficiencia a los voluntarios, a la vez que ofrece una experiencia "
         "transparente y atractiva para la comunidad.", False),
    ])

    add_subhead(doc, "Acceso general (sin registro)")
    add_para(doc, "Cualquier visitante puede:")
    add_bullets(doc, [
        "Consultar la lista de gatos disponibles y acceder a su ficha resumida (foto, nombre, raza, edad, sexo, estado e historia).",
        "Realizar donaciones de forma anónima a través de la pasarela de pago ficticia.",
    ])
    add_para(doc,
        "Al intentar acceder a la sección Colabora, el visitante encuentra un mensaje "
        "“Colabora con nosotros” con opciones para iniciar sesión o registrarse.")

    add_subhead(doc, "Acceso como usuario registrado")
    add_para(doc,
        "Tras registrarse, el usuario accede al detalle completo de cada ficha "
        "felina, que incluye datos técnicos (color, vacunas, carácter), estado de "
        "salud (vacunado, castrado, desparasitado, microchip) y convivencia (apto "
        "para niños, otros gatos, perros). Además puede:")
    add_bullets(doc, [
        "Marcar gatos como favoritos.",
        ("Presentar una solicitud", " de adopción, apadrinamiento o acogida para un gato."),
        ("Consultar e inscribirse", " en las tareas publicadas en Colabora."),
        "Conversar con el asistente virtual con IA, que orienta sobre los gatos disponibles y la asociación.",
        "Realizar donaciones nominales asociadas a su perfil.",
        ("Acceder a su perfil:", " donde encuentra su familia felina (gatos adoptados, acogidos o apadrinados), el estado de sus solicitudes, sus insignias desbloqueadas y un feed con las novedades de los gatos vinculados a él."),
    ])

    add_subhead(doc, "Acceso como administrador")
    add_para(doc,
        "El administrador es el personal de la asociación con responsabilidad "
        "operativa sobre la plataforma. Desde el panel interno puede gestionar las "
        "siguientes áreas:")

    add_h3(doc, "Gestión de fichas felinas")
    add_bullets(doc, [
        "Crear, editar y eliminar fichas.",
        "Gestionar el estado del gato: adoptable, apadrinable o adoptado.",
        "Subir fotos de los animales.",
        "Publicar actualizaciones del gato (título, mensaje y foto), que aparecen en el feed de los usuarios vinculados a ese animal.",
    ])

    add_h3(doc, "Gestión de usuarios y roles")
    add_bullets(doc, [
        "Visualizar la lista de usuarios registrados.",
        "Asignar o revocar el rol de Administrador.",
    ])

    add_h3(doc, "Gestión de solicitudes")
    add_bullets(doc, [
        "Visualizar todas las solicitudes recibidas (adopción y acogida).",
        "Filtrar por estado (pendiente, en revisión, aprobada, rechazada).",
        "Revisar los cuestionarios completos enviados por el usuario.",
        "Cambiar el estado de la solicitud.",
    ])

    add_callout(doc, [
        ("El apadrinamiento sigue un flujo independiente: ", False),
        ("no requiere aprobación", True),
        (" y se gestiona directamente desde el módulo de pago recurrente.", False),
    ], label=None)

    add_h3(doc, "Logística interna (lista de tareas)")
    add_bullets(doc, [
        "Crear tareas, clasificadas por tipo (voluntariado, material o evento) y nivel de urgencia.",
        "Editar tareas existentes.",
        "Cambiar su estado (pendiente, en curso, finalizada).",
        "Asignar un código postal a la tarea para que los usuarios puedan filtrar por proximidad.",
        "Consultar los voluntarios inscritos en cada tarea con sus datos de contacto.",
    ])

    add_subhead(doc, "Solicitud de adopción")
    add_para(doc,
        "Para presentar una solicitud, el usuario completa un cuestionario "
        "obligatorio que incluye:")
    add_bullets(doc, [
        "Experiencia previa con animales.",
        "Motivo de la adopción.",
        "Condiciones de la vivienda.",
    ])
    add_para(doc,
        "Una vez enviada, puede consultar el estado de su solicitud (pendiente, en "
        "revisión, aprobada, rechazada) desde su perfil. Si una solicitud previa fue "
        "rechazada, el usuario puede volver a presentar una nueva para el mismo "
        "gato; en cualquier otro estado, el sistema impide duplicarla.")

    add_subhead(doc, "Tipos de usuarios")
    add_para(doc,
        "El sistema cuenta con un registro simple mediante correo electrónico y "
        "contraseña, que otorga automáticamente el rol inicial de Usuario (USER).")
    add_bullets(doc, [
        ("Visitante anónimo:", " cualquier persona que accede sin registrarse."),
        ("Usuario registrado (USER):", " persona dada de alta en la plataforma con acceso ampliado a las fichas, posibilidad de presentar solicitudes, marcar favoritos, conversar con el asistente virtual y participar en las tareas de la asociación."),
        ("Administrador (ADMIN):", " personal de la asociación con responsabilidad operativa. Acceso total al panel interno, gestión de fichas felinas, gestión de tareas, gestión de usuarios y revisión de solicitudes."),
    ])

    add_subhead(doc, "Gestión de roles")
    add_para(doc,
        "Cualquier persona puede registrarse y obtener automáticamente el rol de "
        "Usuario. La elevación al rol de Administrador se realiza manualmente desde "
        "el panel interno: un administrador existente accede al listado de usuarios "
        "y le asigna el rol. La promoción es siempre una decisión del equipo de la "
        "asociación.")

    add_subhead(doc, "Fichas felinas")
    add_para(doc,
        "Cada gato cuenta con una única ficha en el sistema, aunque la información "
        "mostrada varía según el estado de autenticación del visitante. Los datos "
        "completos incluyen:")
    add_bullets(doc, [
        ("Identificación:", " nombre, raza, color, edad, sexo."),
        ("Su historia:", " descripción narrativa del gato."),
        ("Estado actual:", " adoptable, apadrinable o adoptado."),
        ("Salud:", " vacunado, castrado, desparasitado, microchip y observaciones de vacunas."),
        ("Convivencia:", " apto para niños, apto con otros gatos, apto con perros."),
        ("Carácter:", " descripción del temperamento del animal."),
        ("Diario del gato:", " actualizaciones (foto + título + mensaje) que la asociación publica para mostrar la evolución del animal."),
    ])

    add_subhead(doc, "Pizarra comunitaria")
    add_para(doc,
        "La sección Colabora funciona como una pizarra comunitaria donde la "
        "asociación publica las tareas que necesita cubrir. Solo es accesible para "
        "usuarios registrados.")
    add_para(doc, "Cada tarea publicada muestra:")
    add_bullets(doc, [
        ("Tipo de necesidad:", " voluntariado, material o evento."),
        ("Resumen", " del trabajo a realizar."),
        ("Nivel de urgencia:", " alta, media o baja."),
        ("Estado actual:", " pendiente, en curso o finalizada."),
        ("Código postal de referencia,", " que permite a los usuarios filtrar las tareas por proximidad geográfica."),
    ])
    add_para(doc,
        "Cualquier usuario registrado puede inscribirse en una tarea desde la propia "
        "plataforma. La asociación recibe el aviso, se pone en contacto con el "
        "voluntario por canal privado para coordinar los detalles, y actualiza el "
        "estado de la tarea conforme avanza el trabajo.")

    add_subhead(doc, "Módulo de Apadrinamiento")
    add_para(doc,
        "El apadrinamiento permite a un usuario registrado vincularse económicamente "
        "con un gato concreto a través de una aportación mensual recurrente. El "
        "sistema gestiona:")
    add_bullets(doc, [
        ("Vinculación inicial:", " elección del gato, configuración del importe mensual y de los datos de pago a través de la pasarela ficticia."),
        ("Pago recurrente:", " la plataforma calcula y registra automáticamente los cobros mensuales, mostrando al usuario la fecha del próximo cobro y el historial de aportaciones."),
        ("Área privada del padrino:", " dentro del perfil, acceso a la ficha del gato apadrinado, su historial de pagos y las actualizaciones publicadas por la asociación."),
        ("Modificación o cancelación:", " el usuario puede cambiar el importe mensual o cancelar el apadrinamiento cuando lo desee."),
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Mm(3)
    p.paragraph_format.line_spacing = 1.55
    styled_run(p, "A diferencia de las solicitudes de adopción y acogida, el apadrinamiento ", size=11, color=C_TEXTO)
    styled_run(p, "no requiere aprobación", size=11, color=C_TEXTO, bold=True)
    styled_run(p, ": se activa de forma inmediata al completar el pago.", size=11, color=C_TEXTO)

    add_subhead(doc, "Asistente virtual con IA")
    add_para(doc,
        "La plataforma incluye un asistente conversacional integrado, accesible para "
        "los usuarios registrados, que orienta a quienes se interesan por la "
        "asociación. El asistente conoce el catálogo de gatos disponibles —en "
        "adopción, apadrinamiento o acogida— y responde dudas sobre el funcionamiento "
        "del refugio, los procesos de solicitud o las características de los "
        "animales. Sirve como primer punto de contacto y reduce la carga de atención "
        "del equipo humano para preguntas frecuentes.")

    add_subhead(doc, "Personalización del perfil")
    add_para(doc, "Cada usuario registrado puede personalizar su perfil con:")
    add_bullets(doc, [
        ("Avatar:", " elegir como imagen de perfil la foto de uno de los gatos del refugio."),
        ("Portada:", " seleccionar entre cinco temas visuales (morado, crema, amanecer, lavanda y bosque)."),
        ("Datos personales editables:", " nombre, apellido, teléfono y código postal — este último también se utiliza como referencia de proximidad para las tareas de la pizarra comunitaria."),
    ])

    add_subhead(doc, "Sistema de insignias y reconocimientos")
    add_para(doc,
        "A medida que el usuario se implica con la asociación, desbloquea insignias "
        "visibles en su perfil:")
    add_bullets(doc, [
        ("Familia:", " se concede tras adoptar, acoger o apadrinar al menos un gato."),
        ("Donante:", " se concede tras realizar al menos una donación."),
        ("Veterano:", " se concede al cumplir un año de antigüedad en la plataforma."),
    ])
    add_para(doc,
        "Las insignias funcionan como un reconocimiento simbólico al compromiso del "
        "usuario y refuerzan su vínculo emocional con la asociación.")

    add_close_rule(doc)


def build_seccion_6_1(doc):
    add_section_title(doc, "6.1 Herramientas de desarrollo software utilizadas")

    add_lead(doc, [
        ("El proyecto se ha construido como una ", False),
        ("aplicación cliente–servidor", True),
        (", separando con claridad la capa de presentación (frontend) de la capa de "
         "negocio y persistencia (backend), y comunicándose entre sí mediante una ", False),
        ("API REST sobre HTTP", True),
        (". Las tecnologías y herramientas empleadas se agrupan a continuación según "
         "su función.", False),
    ])

    add_subhead(doc, "Frontend (capa de cliente)")
    add_bullets(doc, [
        ("Angular 21.1:", " framework principal de la interfaz, organizado con el modelo de componentes standalone que introduce esta versión."),
        ("TypeScript 5.9:", " lenguaje principal del frontend, que aporta tipado estático y robustez al código."),
        ("RxJS 7.8:", " librería de programación reactiva para gestionar las llamadas HTTP y la propagación de cambios en los servicios."),
        ("Express 5.1 y @angular/ssr:", " servidor Node y motor de renderizado en servidor (SSR) que permite servir páginas pre-renderizadas y mejora la accesibilidad y la indexación."),
        ("Vitest 4 y jsdom:", " entorno de pruebas unitarias del frontend."),
        ("Font Awesome 6.5:", " iconografía del proyecto, cargada por CDN."),
        ("Google Fonts:", " tipografías Inter, Indie Flower y Cormorant Garamond."),
    ])

    add_subhead(doc, "Backend (capa de servidor)")
    add_bullets(doc, [
        ("Spring Boot 3.4.2:", " framework principal del backend; gestiona la inyección de dependencias, la configuración y el arranque."),
        ("Java 21:", " lenguaje y plataforma de ejecución."),
        ("Maven:", " sistema de gestión de dependencias y construcción del proyecto."),
        ("Spring Web:", " implementación de los controladores REST y modelo MVC."),
        ("Spring Data JPA + Hibernate:", " capa de persistencia y mapeo objeto-relacional."),
        ("Spring Security y JJWT 0.11:", " autenticación y autorización con tokens JSON Web Token."),
        ("Spring Validation:", " validación de los datos de entrada."),
        ("Lombok 1.18:", " reducción de código repetitivo en las entidades."),
        ("MapStruct 1.6:", " generación automática de mappers entre entidades y DTOs."),
    ])

    add_subhead(doc, "Base de datos")
    add_bullets(doc, [
        ("PostgreSQL:", " gestor de base de datos relacional utilizado por la aplicación, ejecutado en entorno local durante el desarrollo."),
        ("H2:", " base de datos en memoria, empleada para pruebas y arranque rápido del entorno de desarrollo."),
    ])

    add_subhead(doc, "Servicios externos")
    add_bullets(doc, [
        ("Cloudinary:", " servicio en la nube para la subida y entrega de imágenes (fotos de gatos, avatares, actualizaciones del diario)."),
        ("Google Gemini API:", " modelo de lenguaje que da soporte al asistente conversacional integrado en la plataforma."),
    ])

    add_subhead(doc, "Control de versiones e integración continua")
    add_bullets(doc, [
        ("Git y GitHub:", " control de versiones del código y alojamiento del repositorio."),
        ("GitHub Actions:", " integración continua que ejecuta los tests del frontend automáticamente en cada push al repositorio."),
    ])

    add_subhead(doc, "Herramientas locales de desarrollo")
    add_bullets(doc, [
        ("Node.js 20+ y npm 10:", " entorno de ejecución y gestor de paquetes del frontend."),
        ("IntelliJ IDEA:", " IDE principal del backend (Java + Spring Boot)."),
        ("WebStorm:", " IDE principal del frontend (Angular + TypeScript)."),
        ("Postman:", " pruebas manuales de los endpoints REST durante el desarrollo."),
    ])

    add_close_rule(doc)


def build_seccion_6_2(doc):
    add_section_title(doc, "6.2 Diseño del proyecto")

    add_lead(doc, [
        ("El diseño de Refugio del Sol se ha desarrollado siguiendo un ", False),
        ("enfoque iterativo", True),
        (": partiendo de un lenguaje visual y una paleta cromática definidos al inicio, "
         "las pantallas se han construido y refinado de forma incremental, validando "
         "cada pieza con la siguiente. Este planteamiento, propio del iterative design, "
         "ha permitido ajustar las decisiones a medida que aparecían restricciones "
         "reales del código y necesidades concretas de los usuarios, en lugar de "
         "comprometerse de antemano con un mockup detallado.", False),
    ])

    add_para(doc,
        "A continuación se describen las decisiones estructurales y visuales más "
        "relevantes que han guiado la construcción del proyecto.")

    add_subhead(doc, "Arquitectura general")
    add_para(doc,
        "La aplicación se ha planteado como una arquitectura cliente–servidor "
        "desacoplada, en la que el frontend (Angular) y el backend (Spring Boot) "
        "son proyectos independientes que se comunican exclusivamente a través de "
        "una API REST. Esta separación aporta:")
    add_bullets(doc, [
        ("Independencia tecnológica:", " cada capa puede evolucionar a su ritmo."),
        ("Escalabilidad:", " ambas partes pueden desplegarse y dimensionarse por separado."),
        ("Mantenibilidad:", " las responsabilidades quedan claras y los equipos podrían dividirse por capa."),
    ])
    add_para(doc,
        "La comunicación se realiza mediante HTTP + JSON, utilizando los verbos REST "
        "estándar (GET, POST, PUT, DELETE) y códigos de estado convencionales para "
        "señalizar errores.")

    add_subhead(doc, "Diseño del backend por capas")
    add_para(doc,
        "El backend sigue el patrón en capas habitual de Spring Boot, garantizando "
        "una separación de responsabilidades clara:")
    add_bullets(doc, [
        ("Controladores REST:", " reciben las peticiones HTTP, validan los parámetros y delegan al servicio. No contienen lógica de negocio."),
        ("Servicios:", " orquestan los casos de uso, aplican las reglas de negocio y coordinan con otros servicios."),
        ("Repositorios:", " acceden a la base de datos a través de Spring Data JPA, abstrayendo el SQL."),
        ("Dominio:", " entidades JPA que representan los objetos del modelo (Gato, Usuario, Solicitud, Tarea, etc.) con sus relaciones."),
    ])
    add_para(doc,
        "La autenticación se gestiona mediante tokens JWT firmados, que el frontend "
        "incluye en cada petición autenticada. El servidor verifica el token y "
        "resuelve los permisos antes de ejecutar la acción.")

    add_subhead(doc, "Diseño del frontend")
    add_para(doc,
        "El frontend es una Single Page Application construida con Angular y "
        "componentes standalone. Las decisiones clave fueron:")
    add_bullets(doc, [
        ("Routing protegido:", " las rutas administrativas (/admin/...) solo son accesibles mediante guards (authGuard y adminGuard)."),
        ("Componentes reutilizables:", " patrones como las tarjetas de gato, los modales de confirmación o el sistema de notificaciones toast se reutilizan en toda la aplicación."),
        ("Servicios inyectables:", " cada entidad tiene su servicio Angular (GatoService, SolicitudService, etc.) que centraliza las llamadas HTTP."),
        ("Renderizado en servidor (SSR):", " habilitado con @angular/ssr, mejora la accesibilidad y la indexación inicial de la web."),
    ])

    add_subhead(doc, "Identidad visual y lenguaje de marca")
    add_para(doc,
        "La marca Refugio del Sol se ha construido sobre un lenguaje visual cálido "
        "que transmite cercanía, esperanza y profesionalidad sin caer en lo infantil:")
    add_bullets(doc, [
        ("Paleta cromática:", " un primario mostaza (#E9C46A) que evoca el sol y la calidez del hogar; un secundario sage (#80AB9F) que aporta serenidad natural; y un texto en berenjena (#4F3D63) que da elegancia sin perder legibilidad. Los fondos se construyen sobre crema (#F2E9E4) y blanco para mantener la sensación de espacio limpio."),
        ("Tipografías:", " Inter para el cuerpo de texto (alta legibilidad en pantalla); Indie Flower para los captions tipo polaroid del diario del gato, que humanizan el tono; y Cormorant Garamond para citas y elementos editoriales."),
        ("Lenguaje plástico:", " bordes redondeados generosos (15 a 30 px), sombras suaves multicapa y tarjetas blancas que destacan sobre los fondos crema. La interfaz se siente táctil, no agresiva."),
        ("Animaciones sutiles:", " efecto Ken Burns en el hero (zoom muy lento), reveals al hacer scroll (fade + translate breve) y micro-rotaciones en las polaroides del diario, restauradas al hacer hover. Las animaciones acompañan la lectura sin distraer."),
    ])

    add_subhead(doc, "Decisiones de UX destacadas")
    add_para(doc,
        "A lo largo del desarrollo se han tomado decisiones de experiencia de "
        "usuario que vale la pena explicitar:")
    add_bullets(doc, [
        ("Acceso parcial sin registro.", " El visitante anónimo ve la lista de gatos y su ficha resumida. La información sensible y las acciones (favorito, solicitudes) requieren registro. Esta gradación da transparencia a quien llega de paso e incentiva el alta a quien quiere implicarse, sin convertir el registro en una barrera de entrada."),
        ("Coordinación por canal privado.", " Cuando un voluntario se inscribe en una tarea, la asociación recibe sus datos de contacto y se coordina con él por teléfono o correo. Esto evita publicar direcciones exactas en una pizarra pública y respeta la privacidad de las casas de acogida."),
        ("Bloqueo de solicitudes duplicadas con feedback claro.", " Si el usuario ya tiene una solicitud activa para un gato, el sistema no le permite enviar otra y le muestra un aviso amable, evitando que su gestión anterior se sobrescriba sin su conocimiento."),
        ("Insignias como gamificación de bajo coste.", " En lugar de un sistema de puntos complejo, las insignias (Familia, Donante, Veterano) reconocen hitos relevantes y refuerzan el vínculo emocional sin sobrecargar la interfaz ni introducir incentivos artificiales."),
        ("Asistente virtual para reducir fricción.", " El chat con IA orienta a los usuarios registrados sobre la asociación y los gatos disponibles, reduciendo la carga de atención del equipo humano para preguntas frecuentes."),
    ])

    add_close_rule(doc)


# =============================================================================
# Construcción del documento
# =============================================================================

def build():
    doc = Document()

    style_normal = doc.styles["Normal"]
    style_normal.font.name = FUENTE
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = C_TEXTO
    rPr = style_normal.element.get_or_add_rPr()
    existing = rPr.find(qn("w:rFonts"))
    rFonts = existing if existing is not None else OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FUENTE)
    rFonts.set(qn("w:hAnsi"), FUENTE)
    if rFonts.getparent() is None:
        rPr.append(rFonts)

    # =========================================================================
    # SECCIÓN 1 — PORTADA
    # =========================================================================
    section1 = doc.sections[0]
    set_section_a4(section1)
    set_section_margins(section1, top=0, bottom=0, left=0, right=0)
    section1.different_first_page_header_footer = False

    # Banda superior mostaza (se mantiene como acento de identidad)
    band = doc.add_table(rows=1, cols=1)
    band.autofit = False
    band.columns[0].width = Mm(210)
    band_cell = band.rows[0].cells[0]
    band_cell.width = Mm(210)
    set_cell_background(band_cell, "E9C46A")
    remove_cell_borders(band_cell)
    band_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    band_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    tr = band.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "454")
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

    p_spacer = doc.add_paragraph()
    reset_horizontal_indent(p_spacer)
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Mm(20)
    styled_run(p_spacer, "", size=1)

    # Eyebrow
    p_eyebrow = doc.add_paragraph()
    reset_horizontal_indent(p_eyebrow)
    p_eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eyebrow.paragraph_format.space_after = Mm(14)
    styled_run(p_eyebrow, "PROYECTO FINAL · 2º DAW",
               size=9, color=C_SAGE_OSC, letter_spacing=4, bold=True)

    # Logo centrado
    if LOGO.exists():
        p_logo = doc.add_paragraph()
        reset_horizontal_indent(p_logo)
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Mm(12)
        run_logo = p_logo.add_run()
        run_logo.add_picture(str(LOGO), width=Mm(60))
    else:
        print(f"[aviso] Logo no encontrado en {LOGO}")

    # Título técnico
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.left_indent = Mm(20)
    p_title.paragraph_format.right_indent = Mm(20)
    p_title.paragraph_format.space_after = Mm(8)
    styled_run(p_title, "Plataforma Web para la Gestión Integral",
               size=26, bold=True, color=C_BERENJENA)
    p_title.add_run().add_break()
    styled_run(p_title, "de Asociaciones Felinas",
               size=26, bold=True, color=C_BERENJENA)

    # Línea sage corta
    p_rule = doc.add_paragraph()
    reset_horizontal_indent(p_rule)
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_after = Mm(8)
    styled_run(p_rule, " " * 20, size=1)
    add_horizontal_line(p_rule, hex_color="80AB9F", size_pt=8)

    # CAMBIO 1: Marca y eslogan claramente centrados, sin indentación residual
    p_brand = doc.add_paragraph()
    reset_horizontal_indent(p_brand)
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_brand.paragraph_format.space_after = Mm(3)
    styled_run(p_brand, "Refugio del Sol", size=22, color=C_MOSTAZA_OSC)

    p_tag = doc.add_paragraph()
    reset_horizontal_indent(p_tag)
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_after = Mm(20)
    styled_run(p_tag, "“Que el sol de su vida sea un hogar”",
               size=14, italic=True, color=C_BERENJENA, font="Cambria")

    # Bloque meta (autora, curso, centro)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.autofit = False
    meta_table.columns[0].width = Mm(45)
    meta_table.columns[1].width = Mm(75)
    meta_data = [
        ("AUTORA", "Viviana Salazar"),
        ("CURSO", "2º Desarrollo de Aplicaciones Web"),
        ("CENTRO", "IES Belén"),
        ("CURSO ACADÉMICO", "2025 / 2026"),
    ]
    for i, (k, v) in enumerate(meta_data):
        c0 = meta_table.rows[i].cells[0]
        c0.width = Mm(45)
        c1 = meta_table.rows[i].cells[1]
        c1.width = Mm(75)
        remove_cell_borders(c0)
        remove_cell_borders(c1)
        set_cell_left_border(c0, "E9C46A", size_pt=20)
        p_k = c0.paragraphs[0]
        p_k.paragraph_format.left_indent = Mm(4)
        styled_run(p_k, k, size=8, bold=True, color=C_SAGE_OSC, letter_spacing=2)
        p_v = c1.paragraphs[0]
        styled_run(p_v, v, size=11, bold=True, color=C_BERENJENA)

    tbl_pr = meta_table._tbl.tblPr
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(int(28 * 56.7)))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    p_filler = doc.add_paragraph()
    p_filler.paragraph_format.space_after = Mm(40)
    styled_run(p_filler, "", size=1)

    # Banda inferior sage
    band2 = doc.add_table(rows=1, cols=1)
    band2.autofit = False
    band2.columns[0].width = Mm(210)
    band2_cell = band2.rows[0].cells[0]
    set_cell_background(band2_cell, "80AB9F")
    remove_cell_borders(band2_cell)
    tr2 = band2.rows[0]._tr
    trPr2 = tr2.get_or_add_trPr()
    trHeight2 = OxmlElement("w:trHeight")
    trHeight2.set(qn("w:val"), "284")
    trHeight2.set(qn("w:hRule"), "exact")
    trPr2.append(trHeight2)

    # =========================================================================
    # SECCIÓN 2 — RESTO DEL DOCUMENTO
    # =========================================================================
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_a4(section2)
    set_section_margins(section2, top=24, bottom=22, left=22, right=22)
    section2.different_first_page_header_footer = False

    section2.header.is_linked_to_previous = False
    section2.footer.is_linked_to_previous = False

    # ---------- Header (CAMBIO 2: sin texto a la derecha y sin línea mostaza) ----------
    header_table = section2.header.add_table(rows=1, cols=2, width=Mm(166))
    header_table.autofit = False
    header_table.columns[0].width = Mm(83)
    header_table.columns[1].width = Mm(83)
    hl, hr = header_table.rows[0].cells
    remove_cell_borders(hl)
    remove_cell_borders(hr)
    p_hl = hl.paragraphs[0]
    styled_run(p_hl, "Refugio del Sol", size=9, color=C_BERENJENA, letter_spacing=0.6)
    add_horizontal_line(p_hl, hex_color=C_LINEA_GRIS, size_pt=2)
    # Celda derecha vacía pero con la misma línea inferior gris para continuidad
    p_hr = hr.paragraphs[0]
    p_hr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_horizontal_line(p_hr, hex_color=C_LINEA_GRIS, size_pt=2)

    # ---------- Footer ----------
    footer_table = section2.footer.add_table(rows=1, cols=2, width=Mm(166))
    footer_table.autofit = False
    footer_table.columns[0].width = Mm(83)
    footer_table.columns[1].width = Mm(83)
    fl, fr = footer_table.rows[0].cells
    remove_cell_borders(fl)
    remove_cell_borders(fr)
    p_fl = fl.paragraphs[0]
    add_top_line(p_fl, hex_color="80AB9F", size_pt=4)
    styled_run(p_fl, "Viviana Salazar  ·  2º DAW", size=9, color=C_BERENJENA)
    p_fr = fr.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_top_line(p_fr, hex_color="80AB9F", size_pt=4)
    styled_run(p_fr, "— ", size=9, color=C_BERENJENA, italic=True)
    add_field(p_fr, "PAGE", color=C_BERENJENA, italic=True)
    styled_run(p_fr, " —", size=9, color=C_BERENJENA, italic=True)

    # =========================================================================
    # ÍNDICE
    # =========================================================================
    p_eyeb = doc.add_paragraph()
    p_eyeb.paragraph_format.space_before = Mm(4)
    p_eyeb.paragraph_format.space_after = Mm(2)
    styled_run(p_eyeb, "TABLA DE CONTENIDOS", size=9, color=C_SAGE_OSC,
               bold=True, letter_spacing=4)

    p_titulo_idx = doc.add_paragraph()
    p_titulo_idx.paragraph_format.space_after = Mm(2)
    styled_run(p_titulo_idx, "Índice", size=34, bold=True, color=C_BERENJENA)

    p_idx_rule = doc.add_paragraph()
    p_idx_rule.paragraph_format.space_after = Mm(10)
    styled_run(p_idx_rule, " " * 12, size=1)
    add_horizontal_line(p_idx_rule, hex_color="E9C46A", size_pt=12)

    INDICE = [
        ("01", "Título del proyecto", "01", []),
        ("02", "Índice", "02", []),
        ("03", "Estudio del sector productivo", "04", []),
        ("04", "Finalidad del proyecto", "07", []),
        ("05", "Descripción básica del proyecto", "11", []),
        ("06", "Descripción detallada del proyecto", "15", [
            ("6.1",  "Herramientas de desarrollo software utilizadas", "16"),
            ("6.2",  "Diseño del proyecto", "19"),
            ("6.3",  "Esquema de funcionamiento de la aplicación", "23"),
            ("6.4",  "Descomposición modular e interrelación entre módulos", "27"),
            ("6.5",  "Descripción de los módulos del proyecto", "31"),
            ("6.6",  "Descripción de la interfaz de usuario", "38"),
            ("6.7",  "Descripción de listados e informes", "44"),
            ("6.8",  "Manual de usuario", "47"),
            ("6.9",  "Manual de instalación o despliegue", "53"),
            ("6.10", "Presentación de la aplicación", "58"),
        ]),
        ("07", "Planificación del proyecto", "61", []),
        ("08", "Control de la ejecución", "65", []),
        ("09", "Dificultades encontradas", "68", []),
        ("10", "Propuestas de mejora", "71", []),
        ("11", "Conclusión final", "74", []),
    ]

    for num, title, page, subs in INDICE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(
            Mm(166), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        styled_run(p, f"{num}    ", size=11, bold=True, color=C_MOSTAZA_OSC)
        styled_run(p, title, size=11, color=C_BERENJENA, bold=True)
        styled_run(p, "\t", size=11)
        styled_run(p, page, size=10, color=C_TEXTO_SUAVE)

        for sn, st, sp in subs:
            ps = doc.add_paragraph()
            ps.paragraph_format.left_indent = Mm(14)
            ps.paragraph_format.space_before = Pt(0)
            ps.paragraph_format.space_after = Pt(0)
            ps.paragraph_format.tab_stops.add_tab_stop(
                Mm(152), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
            )
            styled_run(ps, f"{sn}    ", size=10, color=C_SAGE)
            styled_run(ps, st, size=10, color=C_SAGE_OSC)
            styled_run(ps, "\t", size=10)
            styled_run(ps, sp, size=10, color=C_TEXTO_SUAVE)

    doc.add_page_break()

    # =========================================================================
    # SECCIONES DE CONTENIDO
    # =========================================================================
    build_seccion_3(doc)
    doc.add_page_break()
    build_seccion_4(doc)
    doc.add_page_break()
    build_seccion_5(doc)
    doc.add_page_break()
    build_seccion_6_1(doc)
    doc.add_page_break()
    build_seccion_6_2(doc)

    doc.save(SALIDA)
    print(f"[OK] Documento generado: {SALIDA}")


if __name__ == "__main__":
    build()
